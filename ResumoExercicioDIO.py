# !pip install requests python-docx
#//////////////////////////////////////////////////

#Cria um tradutor de texto:
import requests
from docx import Document
import os

subscription_key = "CHAVE_API_AZURE"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = "eastus2"
language_destination = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type':'application/json',
      'x-ClientTraceId':str(os.urandom(16))
  }
  body = [{
      'text': text
  }]
  params = {
      'api-version':'3.0',
      'from':'es',
      'to':target_language
  }
  request = requests.post(constructed_url, params=params, headers=headers,json=body)
  response = request.json()
  # The key was misspelled, it should be "translations"
  return response[0]["translations"][0]["text"]

#//////////////////////////////////////////////////

translator_text("Solo yo com mi cara de nada", language_destination)

#//////////////////////////////////////////////////

#Pega um documento e traduz
def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    # Armazenando o resultado da tradução em translated_text:
    translated_text = translator_text(paragraph.text, language_destination)
    # Adicionando translated_text à lista full_text:
    full_text.append(translated_text)  # Corrected: append the translated text

  # Criando um novo documento para o conteúdo traduzido:
  translated_doc = Document()
  for line in full_text:
    print(line)
    # Adicionando cada linha traduzida como um parágrafo no novo documento:
    translated_doc.add_paragraph(line)

  # Criando um nome de arquivo para o documento traduzido:
  path_translated = path.replace(".docx", f"_{language_destination}.docx")

  # Salvando o documento traduzido:
  translated_doc.save(path_translated)  # Corrigido para salvar o documento traduzido

  return path_translated

input_file = "/content/CaraDaNada.docx"
translate_document(input_file)

#//////////////////////////////////////////////////
#!pip install requests beautifulsoup4 openai langchain-openai

#//////////////////////////////////////////////////

#Extrai um texto da web
import requests
from bs4 import BeautifulSoup

def extract_text_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
      soup = BeautifulSoup(response.text, 'html.parser')
      for script_or_style in soup(['script', 'style']):
        script_or_style.decompose()
      texto = soup.get_text(separator= ' ')
      #Limpar texto
      linhas = (line.strip() for line in texto.splitlines())
      parts = (phrase.strip() for line in linhas for phrase in line.split("  "))
      texto_limpo = '\n'.join(part for part in parts if part)
      return texto_limpo
    else:
      print(f"Erro ao acessar a URL: {response.status_code}")
      return None
    
    text = soup.get_text()
    return text

extract_text_from_url('https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo')

#//////////////////////////////////////////////////
#!pip install langchain-openai

#//////////////////////////////////////////////////

#Pega um texto da web e traduz
from langchain_openai.chat_models.azure import AzureChatOpenAI # Use the correct import path for AzureChatOpenAI

client = AzureChatOpenAI(
    azure_endpoint="https://openai-trans-bootcamp-eastus-teste.openai.azure.com/",
    api_key="CHAVE_API_AZURE",
    api_version="2023-07-01-preview",  # Check for the most recent version
    deployment_name="gpt-4",
    max_retries=0
)

def translate_article(text, lang):
  messages = [
      ("system", "Você atua como tradutor de textos"),
      ("user", f"Traduza o {text} para o idioma {lang} e responda em markdown")
  ]

  response = client(messages)
  print(response.content)
  return response.content

translate_article("Stay with me", "portugues")

#//////////////////////////////////////////////////

url = 'https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo'
text = extract_text_from_url(url)
article = translate_article(text, "pt-br") 

print(article)

#//////////////////////////////////////////////////